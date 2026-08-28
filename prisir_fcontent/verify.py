# -*- coding: utf-8 -*-
"""Prisir Fcontent 端到端自检:造真 docx/pdf/txt → 索引 → 内容搜索命中 → snippet → clear。

  python verify.py
"""
import os, shutil, sys, tempfile, time

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(HERE))  # 包父目录,import prisir_fcontent
from prisir_fcontent import Fcontent  # noqa: E402

TAG = "__fcontent_verify__"  # 测试标记,末尾自清理


def build_tree(root):
    """造带已知内容真文件:docx(中文)、pdf(英文)、txt(中英混合)。"""
    os.makedirs(root, exist_ok=True)
    paths = {}

    # txt:中英混合(直读)
    p_txt = os.path.join(root, f"{TAG}notes.txt")
    with open(p_txt, "w", encoding="utf-8") as f:
        f.write("季度报告要点:营收增长稳健,See quarterly revenue growth. 市场份额扩大。")
    paths["txt"] = p_txt

    # docx:中文正文(python-docx 写真文件)
    import docx
    p_docx = os.path.join(root, f"{TAG}季度报告.docx")
    d = docx.Document()
    d.add_paragraph("2026 年季度报告:营收增长超出预期,市场份额稳步扩大。")
    d.add_paragraph("下半年规划:深耕企业级客户。")
    d.save(p_docx)
    paths["docx"] = p_docx

    # pdf:英文正文(pypdf 写真文件)
    from pypdf import PdfWriter
    from pypdf.annotations import FreeText  # noqa: F401  (确保 pypdf 可用)
    p_pdf = os.path.join(root, f"{TAG}report.pdf")
    w = PdfWriter()
    w.add_blank_page(width=612, height=792)
    # pypdf 无法直接写文本层;用 reportlab 风格不便,改往 metadata + 追加可见文本流太重。
    # 简化:pdf 侧只验「能建空 pdf 且 extract 不崩」,正文命中主要靠 docx/txt。
    w.add_metadata({"/Title": f"{TAG} quarterly revenue report"})
    with open(p_pdf, "wb") as f:
        w.write(f)
    paths["pdf"] = p_pdf

    # 不支持类型:png(默认跳过;ocr=True 时才识别)
    # 造一张真带字图(rapidocr 可用时验 OCR 链路;不可用则仍是「跳过」)
    p_png = os.path.join(root, f"{TAG}截图.png")
    try:
        from PIL import Image, ImageDraw, ImageFont
        fnt = None
        for fp in (r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf"):
            if os.path.exists(fp):
                fnt = ImageFont.truetype(fp, 30)
                break
        img = Image.new("RGB", (520, 90), "white")
        ImageDraw.Draw(img).text((12, 24), "季度报告市场份额扩大", font=fnt, fill="black")
        img.save(p_png)
    except Exception:  # noqa: BLE001
        with open(p_png, "wb") as f:  # PIL 不可用则退化成假 PNG(仍验「默认跳过」)
            f.write(b"\x89PNG\r\n\x1a\n")
    paths["png"] = p_png
    return paths


def main():
    tmp = tempfile.mkdtemp(prefix="fcontent_verify_")
    db = os.path.join(tmp, "test.db")
    tree = os.path.join(tmp, "tree")
    try:
        paths = build_tree(tree)
        fc = Fcontent(db)
        print("[1] status 未开启:", fc.status())
        assert fc.status()["enabled"] is False, "初始应为未开启"

        r = fc.enable([tree])
        print(f"[2] enable -> {r}")
        assert r.get("ok"), f"enable 失败: {r}"
        st = fc.status()
        # txt + docx 可索引;pdf(空文本层)/png 跳过(默认不开 OCR)
        print("[3] status 已开启:", st)
        assert st["enabled"] is True and st["indexed_count"] >= 2, f"count 错: {st}"
        ocr_avail = st["ocr"]["available"]
        assert isinstance(ocr_avail, bool), "OCR 位应真探测(bool)"
        assert r.get("ocr") is False, "默认 enable 不应开 OCR"
        # 默认不开 OCR:图片不应进索引
        png_hit = [h for h in fc.search("市场份额")["hits"] if h["path"].endswith(".png")]
        assert not png_hit, f"默认(ocr=False)图片不应进索引: {png_hit}"

        res = fc.search("季度报告")
        names = [os.path.basename(h["path"]) for h in res["hits"]]
        print(f"[4] search('季度报告') total={res['total']} -> {names}")
        assert any("季度报告" in n or "notes" in n for n in names), "应命中含'季度报告'的文件"
        snips = [h["snippet"] for h in res["hits"]]
        assert any("**" in s for s in snips), f"应有匹配片段高亮: {snips}"

        res_en = fc.search("revenue")
        names_en = [os.path.basename(h["path"]) for h in res_en["hits"]]
        print(f"[5] search('revenue') -> {names_en}")
        assert any("notes" in n for n in names_en), "应命中含 revenue 的 txt"

        # 中文短词(验 CJK 分词不丢)
        res_zh = fc.search("市场份额")
        names_zh = [os.path.basename(h["path"]) for h in res_zh["hits"]]
        print(f"[6] search('市场份额') -> {names_zh}")
        assert names_zh, "中文短词'市场份额'应命中"

        # 空查询:按 mtime 倒序列全部
        allr = fc.search("", 50)
        print(f"[7] 空查询 total={allr['total']}")
        assert allr["total"] >= 2

        # OCR 开关:rapidocr 可用时,ocr=True 应把带字 png 索引进来
        if ocr_avail:
            r2 = fc.enable([tree], ocr=True)
            print(f"[8] enable(ocr=True) -> {r2}")
            assert r2.get("ok") and r2.get("ocr") is True, f"ocr=True 未生效: {r2}"
            ocr_hits = [h for h in fc.search("市场份额")["hits"] if h["path"].endswith(".png")]
            print(f"[9] search('市场份额') OCR 命中 -> {[os.path.basename(h['path']) for h in ocr_hits]}")
            assert ocr_hits and all(h["is_ocr"] for h in ocr_hits), "ocr=True 应命中 png 且标 is_ocr"
        else:
            print("[8] rapidocr 未安装,跳过 OCR 用例(可用性已探测为 False)")

        fc.disable()
        st2 = fc.status()
        print("[10] disable 后:", st2)
        assert st2["enabled"] is False and st2["indexed_count"] == 0, "clear 后应空"
        assert st2["ocr_on"] is False, "disable 后 ocr_on 应复位"

        # [11] 重启状态持久化:enable 后 roots/ocr_on/last_scan 写库,新实例(模拟重启)读回
        r3 = fc.enable([tree], ocr=ocr_avail)
        fc.close()
        fc2 = Fcontent(db)  # 同库新实例 ≈ 进程重启
        st3 = fc2.status()
        print(f"[11] 重启后状态 -> roots={len(st3['roots'])} ocr_on={st3['ocr_on']} last_scan={st3['last_scan']}")
        assert st3["roots"], "重启后 roots 应从库读回(不失忆)"
        assert st3["ocr_on"] == bool(ocr_avail), "重启后 ocr_on 应持久化"
        assert st3["last_scan"] > 0, "重启后 last_scan 应持久化"
        fc2.disable()  # 清掉,不留测试索引
        fc2.close()

        # [12] 截图存档:save_shot 落盘 + 元数据 + shot_image 白名单(纯函数级,不走 HTTP)
        import oiagent_web as W  # noqa: PLC0415
        import base64 as _b64, io as _io
        from PIL import Image as _Img, ImageDraw as _ID, ImageFont as _IF
        _fnt = _IF.truetype(r"C:\Windows\Fonts\msyh.ttc", 28)
        _im = _Img.new("RGB", (480, 80), "white")
        _ID.Draw(_im).text((10, 22), "截图存档自检", font=_fnt, fill="black")
        _buf = _io.BytesIO(); _im.save(_buf, "PNG")
        _du = "data:image/png;base64," + _b64.b64encode(_buf.getvalue()).decode()
        code, resp = W._save_shot({"data_url": _du, "page_url": "https://x.example/p",
                                   "page_title": "自检页", "scroll": {"x": 0, "y": 100}, "ts": 1787373900000})
        print(f"[12] save_shot -> HTTP {code} saved={resp.get('saved')} indexed={resp.get('indexed')}")
        assert code == 200 and resp.get("saved") and os.path.isfile(resp["path"])
        meta = W._shot_lookup(resp["path"])
        assert meta and meta["page_url"] == "https://x.example/p" and meta["scroll"]["y"] == 100, f"元数据错: {meta}"
        assert W._shot_in_dir(resp["path"]) is True, "截图目录内应过白名单"
        assert W._shot_in_dir(os.path.abspath(__file__)) is False, "目录外应被白名单拒"
        code2, resp2 = W._save_shot({"data_url": "data:text/html;base64,PGI+"})
        assert code2 == 400, f"非 png 应 400: {code2}"
        try:
            os.remove(resp["path"])  # 清理截图存档目录里的自检图
        except OSError:
            pass
        print("[12] save_shot/元数据/白名单/越权 全过")

        # [13] 叠加翻译管线:overlay_translate(mock 翻译)—— 产物 *.translated.png + 块结构 + 原图不动。
        # 真 LLM 翻译链已人工 E2E 验证(英→中);自检用 mock 保证确定性、离线可跑。
        from prisir_fcontent import overlay_translate as _OVT  # noqa: PLC0415
        if ocr_avail:
            _im2 = _Img.new("RGB", (560, 200), "white")
            _d2 = _ID.Draw(_im2)
            _ef = _IF.truetype(r"C:\Windows\Fonts\arial.ttf", 30)
            _d2.rounded_rectangle([20, 30, 400, 110], radius=14, fill="white", outline="black", width=3)
            _d2.text((45, 52), "Wait... what?", font=_ef, fill="black")
            _sp = os.path.join(tmp, "ov_en.png"); _im2.save(_sp)
            _r = _OVT.overlay_translate(_sp, lambda t: "【译】" + t, dst="zh")
            print(f"[13] overlay_translate -> ok={_r.get('ok')} blocks={len(_r.get('blocks', []))}")
            assert _r.get("ok"), f"叠加翻译失败: {_r.get('error')}"
            assert os.path.isfile(_r["out"]) and _r["out"].endswith(".translated.png"), "产物应 *.translated.png"
            assert _r["blocks"] and _r["blocks"][0]["dst"].startswith("【译】"), "块译文应来自 translate_fn"
            assert os.path.getsize(_sp) and os.path.isfile(_sp), "原图应保留不动"
            # 低质/无字图应 no_text(门控兜底)
            _blank = os.path.join(tmp, "ov_blank.png")
            _Img.new("RGB", (300, 120), "white").save(_blank)
            _rb = _OVT.overlay_translate(_blank, lambda t: "x", dst="zh")
            assert _rb.get("ok") is False and _rb.get("error") in ("no_text", "translate_empty"), \
                f"无字图应 no_text: {_rb}"
            print("[13] 叠加翻译(产物/块结构/原图不动/无字门控)全过")
        else:
            print("[13] OCR 不可用,跳过叠加翻译自检")

        # [14] 抹字版翻译管线:overlay_translate_erase(mock 翻译)—— 真抹字 + 方向可选 + 原图不动。
        # 真 LLM/google_gtx 翻译链已人工 E2E 验证(漫画中→日横排/竖排);自检用 mock 保证确定性、离线可跑。
        if ocr_avail:
            try:
                import numpy, cv2  # noqa: F401
                _cv2_ok = True
            except Exception:  # noqa: BLE001
                _cv2_ok = False
            if not _cv2_ok:
                print("[14] cv2/numpy 不可用,跳过抹字版自检(需 pip install opencv-python numpy)")
            else:
                # 横排英文气泡 → direction="h" 横排译文
                _im3 = _Img.new("RGB", (560, 200), "white")
                _d3 = _ID.Draw(_im3)
                _d3.rounded_rectangle([20, 30, 400, 110], radius=14, fill="white", outline="black", width=3)
                _d3.text((45, 52), "Wait... what?", font=_ef, fill="black")
                _sp3 = os.path.join(tmp, "er_en.png"); _im3.save(_sp3)
                _r3 = _OVT.overlay_translate_erase(_sp3, lambda t, d: "【译】" + t, dst="zh", direction="h")
                print(f"[14] erase(h) -> ok={_r3.get('ok')} blocks={len(_r3.get('blocks', []))}")
                assert _r3.get("ok"), f"抹字横排失败: {_r3.get('error')} {_r3.get('hint')}"
                assert os.path.isfile(_r3["out"]) and _r3["out"].endswith(".translated.png"), "产物应 *.translated.png"
                assert _r3["blocks"] and _r3["blocks"][0]["dst"].startswith("【译】"), "块译文应来自 translate_fn"
                assert _r3["blocks"][0]["vertical"] is False, "direction=h 应横排"
                assert os.path.isfile(_sp3), "原图应保留不动"
                # direction="v" 强制竖排:同一张图应判 vertical=True
                _r4 = _OVT.overlay_translate_erase(_sp3, lambda t, d: "【译】" + t, dst="zh", direction="v")
                assert _r4.get("ok") and _r4["blocks"][0]["vertical"] is True, "direction=v 应竖排"
                # auto:横排英文行形 → 应判横排
                _r5 = _OVT.overlay_translate_erase(_sp3, lambda t, d: "【译】" + t, dst="zh", direction="auto", src_lang="en")
                assert _r5.get("ok") and _r5["blocks"][0]["vertical"] is False, "auto 横排英文应判横排"
                # 无字图应 no_text
                _blank2 = os.path.join(tmp, "er_blank.png")
                _Img.new("RGB", (300, 120), "white").save(_blank2)
                _rb2 = _OVT.overlay_translate_erase(_blank2, lambda t, d: "x", dst="zh")
                assert _rb2.get("ok") is False and _rb2.get("error") in ("no_text", "translate_empty"), \
                    f"无字图应 no_text: {_rb2}"
                print("[14] 抹字版(产物/方向 h·v·auto/原图不动/无字门控)全过")
        print("\n[PASS] fcontent 端到端自检全部通过")
        return 0
    except AssertionError as e:
        print(f"\n[FAIL] {e}")
        return 1
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    sys.exit(main())
